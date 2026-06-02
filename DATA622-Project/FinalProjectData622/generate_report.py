"""
Generate Professional Final Project Report - Word Document
Intelligent Pharmacy Inventory Management System
DATA 622 Machine Learning Final Project
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def set_cell_shading(cell, color_hex):
    """Set background shading for a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_formatted_table(doc, headers, rows, header_color="1F4E79"):
    """Add a professionally formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, header_color)

    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, value in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(value))
            run.font.size = Pt(9)
            if row_idx % 2 == 0:
                set_cell_shading(cell, "F2F7FC")

    return table


def add_formula(doc, text):
    """Add a centered formula paragraph."""
    f = doc.add_paragraph()
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = f.add_run(text)
    run.font.name = 'Cambria Math'
    run.italic = True
    run.font.size = Pt(11)
    return f


def add_caption(doc, text):
    """Add a centered italic caption."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_figure(doc, img_path, caption_text, base_dir="c:/Projects/DATA622/FinalProjectData622"):
    """Add a figure with caption if the image exists."""
    full_path = os.path.join(base_dir, img_path)
    if os.path.exists(full_path):
        doc.add_picture(full_path, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, caption_text)
    doc.add_paragraph()


def create_report():
    """Generate the complete final project report in narrative whitepaper style."""
    doc = Document()

    # ==================== PAGE SETUP ====================
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ==================== STYLES ====================
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15

    for level, (size, color) in {
        'Heading 1': (16, RGBColor(0x1F, 0x4E, 0x79)),
        'Heading 2': (13, RGBColor(0x2E, 0x75, 0xB6)),
        'Heading 3': (11, RGBColor(0x2E, 0x75, 0xB6)),
    }.items():
        h = doc.styles[level]
        h.font.name = 'Calibri'
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = color

    # ==================== TITLE PAGE ====================
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Intelligent Pharmacy Inventory Management")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Minimizing Waste and Preventing Stockouts\nUsing Time-Series Forecasting")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    doc.add_paragraph()
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("_" * 60)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    doc.add_paragraph()

    for text, size, bold in [
        ("Umais Siddiqui, Inna Yedzinovich, David Greer, and Joao De Oliveira", 13, True),
        ("", 12, False),
        ("DATA 622: Machine Learning\nFinal Project", 12, False),
        ("", 12, False),
        ("May 15, 2026", 12, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if text:
            run = p.add_run(text)
            run.font.size = Pt(size)
            run.bold = bold

    doc.add_page_break()

    # ==================== 1. PROBLEM DEFINITION ====================
    doc.add_heading('1. Problem Definition', level=1)

    doc.add_paragraph(
        "Community pharmacies face a unique inventory optimization challenge because medication "
        "shortages and overstocking both produce serious consequences. When a pharmacy under-stocks "
        "critical medications, patients experience dangerous delays in treatment, the pharmacy incurs "
        "emergency reorder surcharges, and customers may switch to competitors. Conversely, "
        "overstocking leads to financial waste when medications expire on the shelf before being "
        "dispensed. This tension, which we call the Optimization Paradox, is especially difficult "
        "because the costs are fundamentally asymmetric: a single stockout event can cause patient "
        "harm and lost lifetime customer value, while the per-unit cost of carrying excess inventory "
        "is comparatively small."
    )

    doc.add_paragraph(
        "Many pharmacies manage this problem with reactive methods such as manual shelf checks, "
        "static min/max reorder rules, or simple historical averages. These approaches fail to "
        "account for the rich temporal structure in pharmacy demand, including weekly cycles where "
        "Sunday demand drops to 30% of weekday baseline, yearly seasonality driven by flu season "
        "and allergy season, holiday effects that create pre-holiday surges followed by closure-day "
        "drops, and long-term prescription volume growth of approximately 11% annually. This project "
        "addresses the problem by building a machine learning system that forecasts demand with "
        "temporal pattern recognition and converts those forecasts into cost-optimal reorder "
        "recommendations using the classical Newsvendor framework from operations research."
    )

    # ==================== 2. SYSTEM DESIGN ====================
    doc.add_heading('2. System Design', level=1)

    doc.add_paragraph(
        "The system follows a modular pipeline architecture organized into five layers. The "
        "data layer generates or ingests pharmacy sales records using a synthetic data engine "
        "calibrated with negative binomial count-data noise, monthly seasonal profiles, weekly "
        "demand patterns, and US holiday calendars. This engine produces 6,304 records covering "
        "eight medications across 788 days. The processing layer prepares data for Prophet "
        "modeling by performing time-based train/test splits that preserve temporal ordering, "
        "with the most recent 60 days held out for evaluation."
    )

    pipeline = doc.add_paragraph()
    pipeline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pipeline.add_run(
        "Synthetic Data \u2192 Data Preparation \u2192 Forecasting Model \u2192 "
        "Decision Engine \u2192 Dashboard"
    )
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_paragraph()

    doc.add_paragraph(
        "The modeling layer trains per-drug Facebook Prophet models with individually tuned "
        "hyperparameters. The decision engine translates 30-day demand forecasts into inventory "
        "actions by computing cost-optimal order quantities, simulating inventory depletion with "
        "FIFO expiration tracking, and assigning urgency levels. Finally, the dashboard layer "
        "provides an interactive Streamlit interface with authentication, role-based access, "
        "and the ability to upload real pharmacy data."
    )

    doc.add_paragraph(
        "The synthetic demand model generates daily sales as D(t) = base \u00D7 S(t) \u00D7 W(t) "
        "\u00D7 (1+g\u00B7t) \u00D7 H(t) \u00D7 \u03B5(t), where base is the drug-specific "
        "mean demand (8\u201330 units/day), S(t) encodes monthly seasonality (multipliers ranging "
        "from 0.70 to 1.70), W(t) captures the day-of-week pattern, g = 0.03% daily growth, "
        "H(t) applies holiday effects, and \u03B5(t) is drawn from a negative binomial distribution "
        "to produce realistic overdispersed count data."
    )

    doc.add_heading('Drug Catalog', level=2)

    drug_headers = ["Drug", "Category", "Base\n(units/day)", "Unit Cost", "Shelf Life", "Seasonal Profile"]
    drug_rows = [
        ["Amoxicillin 500mg", "Antibiotic", "18", "$0.45", "365 d", "Winter spike"],
        ["Metformin 500mg", "Diabetes", "30", "$0.12", "730 d", "Stable"],
        ["Lisinopril 10mg", "Blood Pressure", "25", "$0.08", "730 d", "Stable"],
        ["Albuterol Inhaler", "Respiratory", "8", "$25.00", "365 d", "Winter spike"],
        ["Cetirizine 10mg", "Allergy", "12", "$0.15", "730 d", "Spring spike"],
        ["Azithromycin 250mg", "Antibiotic", "10", "$1.20", "365 d", "Winter spike"],
        ["Omeprazole 20mg", "GI", "20", "$0.10", "730 d", "Holiday spike"],
        ["Sertraline 50mg", "Mental Health", "15", "$0.18", "730 d", "Winter SAD"],
    ]
    add_formatted_table(doc, drug_headers, drug_rows)
    add_caption(doc, "Table 1. Drug catalog with demand characteristics and cost parameters.")

    doc.add_page_break()

    # ==================== 3. MACHINE LEARNING COMPONENTS ====================
    doc.add_heading('3. Machine Learning Components', level=1)

    doc.add_paragraph(
        "The forecasting engine uses Facebook Prophet, an additive regression model that "
        "decomposes time-series into trend, seasonality, and holiday components. Prophet was "
        "selected because it naturally handles multiple seasonalities, produces prediction "
        "intervals that quantify forecast uncertainty, and allows domain knowledge to be encoded "
        "through prior scales. Crucially, Prophet generates posterior predictive samples that "
        "enable non-parametric uncertainty quantification\u2014a feature we exploit in the "
        "Newsvendor optimization described in Section 4."
    )

    doc.add_paragraph("Prophet models each drug's demand as an additive decomposition:")
    add_formula(doc, "y(t) = g(t) + s(t) + h(t) + \u03B5(t)")

    doc.add_paragraph(
        "The trend g(t) is a piecewise linear function with automatic changepoint detection, "
        "controlled by a Laplace prior \u03B4_j ~ Laplace(0, \u03C4) where \u03C4 is the "
        "changepoint_prior_scale. A small \u03C4 (0.01 for Metformin) produces a smooth trend "
        "for chronic medications, while a larger \u03C4 (0.15 for Cetirizine) allows structural "
        "shifts. The seasonality s(t) is modeled via Fourier series with three components: "
        "weekly (period 7, capturing the pronounced Sunday demand trough), yearly (period 365.25, "
        "modeling flu and allergy seasons), and monthly (period 30.5, Fourier order 5, capturing "
        "refill patterns). The seasonality_prior_scale ranges from 5.0 for stable chronic drugs "
        "to 20.0 for highly seasonal medications like Cetirizine."
    )

    hp_headers = ["Drug", "Changepoint\nPrior", "Seasonality\nPrior", "Holidays\nPrior", "Rationale"]
    hp_rows = [
        ["Amoxicillin", "0.10", "15.0", "15.0", "Moderate flexibility; strong flu pattern"],
        ["Metformin", "0.01", "5.0", "5.0", "Very smooth; minimal seasonality"],
        ["Lisinopril", "0.01", "5.0", "5.0", "Stable chronic; predictable refills"],
        ["Albuterol", "0.10", "15.0", "10.0", "Winter respiratory spikes"],
        ["Cetirizine", "0.15", "20.0", "5.0", "Strong spring spike; most seasonal"],
        ["Azithromycin", "0.10", "15.0", "10.0", "Seasonal antibiotic"],
        ["Omeprazole", "0.05", "12.0", "15.0", "Holiday GI stress effects"],
        ["Sertraline", "0.05", "10.0", "5.0", "SAD winter pattern"],
    ]
    add_formatted_table(doc, hp_headers, hp_rows)
    add_caption(doc, "Table 2. Per-drug Prophet hyperparameters with clinical rationale.")

    doc.add_page_break()

    # ==================== 4. COST OPTIMIZATION FRAMEWORK ====================
    doc.add_heading('4. Cost Optimization Framework', level=1)

    doc.add_paragraph(
        "The decision engine is grounded in the Newsvendor model, a classical operations research "
        "framework for single-period inventory decisions under demand uncertainty. The key insight "
        "is that stockout and wastage costs are fundamentally asymmetric. We define per-unit "
        "wastage cost as C_w = c_unit \u00D7 p_exp + c_hold, where p_exp is the expiration "
        "probability (5% for short-shelf-life drugs, 2% for long-shelf-life) and c_hold = $0.03/unit/day. "
        "The per-unit stockout cost is C_s = \u03B1 \u00D7 c_unit + c_emergency + c_churn, where "
        "\u03B1 = 10 is an asymmetric penalty, c_emergency = $1.50/unit, and c_churn = $5.00/unit. "
        "For Amoxicillin, this yields C_s/C_w \u2248 210, meaning understocking is over 200 times "
        "more costly per unit than overstocking."
    )

    doc.add_heading('Newsvendor Critical Ratio', level=2)

    doc.add_paragraph(
        "The Newsvendor model minimizes expected asymmetric loss L(Q) = c_w \u00B7 E[max(Q-D,0)] "
        "+ c_s \u00B7 E[max(D-Q,0)] by ordering at the critical ratio quantile of the demand "
        "distribution:"
    )
    add_formula(doc, "Q* = F\u207b\u00b9(c_s / (c_s + c_w))")

    doc.add_paragraph(
        "where F\u207b\u00b9 is the quantile function of demand. When c_s >> c_w, the critical "
        "ratio approaches 1, pushing the order quantity toward the upper tail of demand. For our "
        "pharmacy system, critical ratios range from 0.995 to 0.996 across drugs, reflecting the "
        "strong prioritization of stockout prevention."
    )

    doc.add_heading('Posterior Predictive Sampling', level=2)

    doc.add_paragraph(
        "A common approach is to assume demand follows a Gaussian distribution and use the "
        "inverse normal CDF to evaluate the critical ratio. However, this assumption can be poor "
        "for pharmacy demand, which is count-valued, often skewed, and exhibits overdispersion. "
        "Instead, we use Prophet's posterior predictive sampling to obtain the demand distribution "
        "non-parametrically. After fitting each Prophet model, we call predictive_samples() to "
        "generate 1,000 posterior draws of future demand for each forecast day. These draws "
        "capture uncertainty in the trend, seasonality components, and observation noise "
        "simultaneously."
    )

    doc.add_paragraph(
        "To find Q*, we sum each posterior draw across the 30-day forecast horizon, producing "
        "1,000 samples of total 30-day demand. The optimal order quantity is then simply the "
        "empirical quantile at the critical ratio:"
    )
    add_formula(doc, "Q* = quantile(D\u0302\u2081, D\u0302\u2082, ..., D\u0302\u2081\u2080\u2080\u2080, CR)")

    doc.add_paragraph(
        "This approach is distribution-free and automatically accounts for skewness, heavy tails, "
        "and any non-Gaussian features of the demand distribution. The expected asymmetric loss "
        "is computed empirically as the sample mean of c_w \u00D7 max(Q*-D\u0302,0) + c_s \u00D7 "
        "max(D\u0302-Q*,0) over all draws. When posterior samples are unavailable (e.g., with the "
        "Holt-Winters fallback), the system reverts to the Gaussian approximation."
    )

    doc.add_heading('Reorder Logic', level=2)

    doc.add_paragraph(
        "The reorder point ROP = (lead_time + safety_stock_days) \u00D7 avg_daily_demand "
        "determines when to trigger an order. The system classifies urgency based on days of "
        "stock remaining: CRITICAL when stock will deplete before a new order can arrive, HIGH "
        "when stock falls below the safety buffer, MEDIUM when coverage drops below 14 days, "
        "and LOW otherwise. The decision engine also simulates inventory over 30 days using "
        "a FIFO policy with shelf-life tracking, where batches age daily and expired units are "
        "discarded, providing realistic waste and stockout projections."
    )

    doc.add_page_break()

    # ==================== 5. APPLICATION DEMONSTRATION ====================
    doc.add_heading('5. Application Demonstration', level=1)

    doc.add_paragraph(
        "The system produces exploratory data analysis visualizations, model evaluation plots, "
        "and decision engine outputs. The following figures illustrate key outputs from the pipeline."
    )

    add_figure(doc, "outputs/eda_daily_sales.png",
               "Figure 1. Daily sales time-series for all eight medications, showing seasonal patterns, "
               "trend growth, and weekend demand drops across the 788-day observation period.")

    add_figure(doc, "outputs/eda_monthly_heatmap.png",
               "Figure 2. Monthly sales heatmap revealing seasonal demand concentration: antibiotics "
               "peak in winter, allergy medications peak in spring.")

    add_figure(doc, "outputs/eval_actual_vs_predicted.png",
               "Figure 3. Actual versus predicted demand on the 60-day test set with 80% prediction intervals.")

    add_figure(doc, "outputs/eval_components_amoxicillin.png",
               "Figure 4. Prophet decomposition for Amoxicillin showing trend, weekly seasonality "
               "(Sunday trough), and yearly seasonality (winter flu peak).")

    add_figure(doc, "outputs/eval_mape_summary.png",
               "Figure 5. Weighted MAPE by drug. Dashed line = 20% target. Stable chronic medications "
               "meet the target; seasonal drugs exceed it.")

    add_figure(doc, "outputs/inventory_simulation.png",
               "Figure 6. Thirty-day inventory simulation with projected stock depletion and reorder point lines.")

    add_figure(doc, "outputs/cost_analysis.png",
               "Figure 7. Total cost of ownership breakdown (left) and service level vs. waste rate (right).")

    doc.add_page_break()

    # ==================== 6. RESULTS AND EVALUATION ====================
    doc.add_heading('6. Results and Evaluation', level=1)

    doc.add_heading('Forecast Accuracy', level=2)

    doc.add_paragraph(
        "The system was evaluated on a 60-day held-out test set using weighted MAPE (wMAPE) as "
        "the primary metric, with a target threshold of 20%. The portfolio-average wMAPE is 23.65%, "
        "with three of eight drugs meeting the target. Stable chronic medications perform best: "
        "Metformin achieves 17.63% wMAPE and Lisinopril 19.71%, reflecting their smooth, "
        "refill-driven demand patterns. Omeprazole also meets the target at 18.84%. In contrast, "
        "highly seasonal drugs are harder to forecast: Cetirizine (spring allergy spike) reaches "
        "34.95% wMAPE, and Albuterol (winter respiratory spike with only 8 units/day base demand) "
        "shows 27.92%. The low base volume of Albuterol amplifies percentage errors; its standard "
        "MAPE of 50.81% is misleading, while the volume-weighted metric provides a fairer picture."
    )

    eval_headers = ["Drug", "wMAPE", "MAPE", "sMAPE", "MAE", "RMSE", "80% CI\nCoverage", "Target\nMet"]
    eval_rows = [
        ["Amoxicillin 500mg", "20.00%", "26.45%", "23.06%", "5.56", "6.69", "65.0%", "Marginal"],
        ["Metformin 500mg", "17.63%", "25.67%", "21.45%", "5.28", "6.13", "60.0%", "Yes"],
        ["Lisinopril 10mg", "19.71%", "25.89%", "22.02%", "5.23", "6.45", "63.3%", "Yes"],
        ["Albuterol Inhaler", "27.92%", "50.81%", "33.19%", "3.60", "4.26", "55.0%", "No"],
        ["Cetirizine 10mg", "34.95%", "43.16%", "43.22%", "3.40", "4.05", "70.0%", "No"],
        ["Azithromycin 250mg", "27.26%", "34.49%", "30.71%", "4.22", "5.35", "56.7%", "No"],
        ["Omeprazole 20mg", "18.84%", "19.59%", "19.32%", "4.48", "6.22", "73.3%", "Yes"],
        ["Sertraline 50mg", "22.92%", "28.55%", "25.15%", "4.44", "5.79", "68.3%", "No"],
    ]
    add_formatted_table(doc, eval_headers, eval_rows)
    add_caption(doc, "Table 3. Model evaluation metrics on 60-day held-out test set. Target: wMAPE < 20%.")

    doc.add_paragraph()

    doc.add_paragraph(
        "Prediction interval coverage ranges from 55% to 73%, falling below the nominal 80% "
        "target. This indicates the intervals are somewhat overconfident. Prophet's expanding-window "
        "cross-validation (12 rolling cutoffs, 30-day horizon) confirms these patterns: Metformin "
        "is most stable at 22.46% CV MAPE, while Albuterol is least stable at 77.63%. Cross-validation "
        "coverage rates of 70\u201378% are closer to the target, suggesting the test period may have "
        "contained atypical demand patterns."
    )

    doc.add_heading('Reorder Recommendations', level=2)

    doc.add_paragraph(
        "The decision engine generates reorder recommendations using the posterior predictive "
        "Newsvendor optimization. Sertraline receives the highest urgency (ORDER SOON/HIGH) "
        "because its estimated stock covers only 10.6 days against a 12-day lead-time-plus-safety "
        "requirement. Cetirizine is the only drug at MONITOR status with 19.6 days of remaining stock, "
        "reflecting lower current demand as allergy season has not yet peaked. Albuterol drives the "
        "highest order cost ($9,750) due to its $25/inhaler unit cost despite moderate volume."
    )

    rec_headers = ["Drug", "Action", "Urgency", "Order\nQty", "Order\nCost", "30-Day\nDemand",
                   "Days\nLeft", "Service\nLevel"]
    rec_rows = [
        ["Amoxicillin 500mg", "PLAN ORDER", "MEDIUM", "800", "$360", "843", "9.4", "31.2%"],
        ["Metformin 500mg", "PLAN ORDER", "MEDIUM", "840", "$101", "946", "12.9", "43.0%"],
        ["Lisinopril 10mg", "PLAN ORDER", "MEDIUM", "720", "$58", "795", "12.5", "41.6%"],
        ["Albuterol Inhaler", "PLAN ORDER", "MEDIUM", "390", "$9,750", "393", "8.7", "29.0%"],
        ["Cetirizine 10mg", "MONITOR", "LOW", "190", "$29", "257", "19.6", "65.4%"],
        ["Azithromycin 250mg", "PLAN ORDER", "MEDIUM", "440", "$528", "458", "9.4", "31.2%"],
        ["Omeprazole 20mg", "PLAN ORDER", "MEDIUM", "670", "$67", "704", "11.5", "38.5%"],
        ["Sertraline 50mg", "ORDER SOON", "HIGH", "620", "$112", "575", "10.6", "35.5%"],
    ]
    add_formatted_table(doc, rec_headers, rec_rows)
    add_caption(doc, "Table 4. Reorder recommendations from the posterior predictive Newsvendor engine.")

    doc.add_paragraph()

    doc.add_heading('Cost Analysis', level=2)

    doc.add_paragraph(
        "The total cost of ownership analysis confirms the fundamental asymmetry in pharmacy "
        "inventory economics. Across the portfolio, understocking costs ($98,476) dominate at "
        "90% of total cost, while wastage costs ($201) represent less than 1%. This validates "
        "the asymmetric penalty structure and demonstrates the system correctly prioritizes "
        "stockout prevention. Holding costs are negligible ($27), suggesting that carrying "
        "additional safety stock is economically rational. Service levels range from 29% to "
        "65%, all below the 95% target, indicating that even with the Newsvendor optimization "
        "the current estimated inventory levels are insufficient for the full 30-day horizon "
        "and prompt reordering is necessary."
    )

    doc.add_page_break()

    # ==================== 7. STRENGTHS, LIMITATIONS, FUTURE WORK ====================
    doc.add_heading('7. Discussion', level=1)

    doc.add_heading('Strengths', level=2)
    doc.add_paragraph(
        "This project addresses a genuine healthcare operations problem and demonstrates the "
        "complete data science workflow from data generation through deployment. The decision "
        "engine is grounded in established operations research theory rather than ad-hoc rules, "
        "with cost functions derived from first principles. By using Prophet's posterior predictive "
        "samples for the Newsvendor optimization rather than assuming Gaussian demand, the system "
        "avoids a distributional assumption that could be quite poor for count-valued, overdispersed "
        "pharmacy demand. Per-drug hyperparameter tuning based on clinical knowledge ensures that "
        "each medication's unique demand dynamics are captured, and the Streamlit dashboard makes "
        "the outputs accessible to non-technical pharmacists."
    )

    doc.add_heading('Limitations', level=2)
    doc.add_paragraph(
        "The system is currently demonstrated on synthetic data, which, despite careful calibration, "
        "cannot fully reproduce the messiness of real pharmacy operations including supplier shortages, "
        "insurance formulary changes, and local prescriber behavior. Forecast accuracy for highly "
        "seasonal drugs (Cetirizine at 35% wMAPE, Albuterol at 28%) exceeds the 20% target, reflecting "
        "irreducible year-to-year variability in disease season intensity. The cost parameters "
        "(\u03B1=10, c_emergency=$1.50, c_churn=$5.00) are fixed assumptions that would vary in "
        "practice by drug, location, and competitive context. The model also does not incorporate "
        "external signals such as CDC flu surveillance data or pollen counts that could improve "
        "seasonal predictions, and the prediction interval coverage of 55\u201373% suggests the "
        "uncertainty estimates may be somewhat overconfident."
    )

    doc.add_heading('Future Improvements', level=2)
    doc.add_paragraph(
        "The most important next step is validating the system on real pharmacy sales data to "
        "test whether the synthetic data assumptions hold. Additional improvements include "
        "integrating external data sources (CDC FluView, pollen counts, weather) as Prophet "
        "regressors, exploring ensemble approaches that combine Prophet with gradient-boosted or "
        "neural forecasting models for difficult drugs, extending the decision engine to handle "
        "multiple suppliers with varying prices and delivery times, allowing pharmacists to edit "
        "cost parameters directly in the dashboard, and applying conformal prediction to recalibrate "
        "the prediction intervals for more reliable uncertainty estimates."
    )

    doc.add_page_break()

    # ==================== 8. CONCLUSION ====================
    doc.add_heading('8. Conclusion', level=1)

    doc.add_paragraph(
        "This project demonstrates how machine learning can be applied to a practical healthcare "
        "operations problem with meaningful real-world impact. By combining Facebook Prophet "
        "time-series forecasting with Newsvendor optimization theory, the system bridges the gap "
        "between statistical prediction and operational action. The use of posterior predictive "
        "sampling for the Newsvendor critical ratio avoids the Gaussian assumption and provides "
        "a distribution-free approach to cost-optimal ordering that naturally accounts for the "
        "non-Gaussian structure of pharmacy demand."
    )

    doc.add_paragraph(
        "The results reveal that chronic medications with stable refill patterns are highly "
        "predictable (wMAPE 17\u201320%), while seasonal drugs remain challenging (wMAPE 28\u201335%). "
        "The cost analysis confirms that understocking dominates the total cost of ownership, "
        "validating the system's asymmetric penalty structure. The most important lesson is that "
        "machine learning is most valuable when connected to a business workflow: a forecast alone "
        "is informative, but a forecast connected to cost optimization, inventory simulation, "
        "and an interactive dashboard becomes a decision-support system that can help pharmacies "
        "reduce waste, improve efficiency, and protect patient access to critical medications."
    )

    doc.add_page_break()

    # ==================== REFERENCES ====================
    doc.add_heading('References', level=1)

    references = [
        "Taylor, S.J. and Letham, B. (2018). Forecasting at Scale. The American Statistician, 72(1), 37-45.",
        "Arrow, K.J., Harris, T., and Marschak, J. (1951). Optimal Inventory Policy. Econometrica, 19(3), 250-272.",
        "Silver, E.A., Pyke, D.F., and Thomas, D.J. (2017). Inventory and Production Management in Supply Chains. 4th Ed., CRC Press.",
        "Centers for Disease Control and Prevention. FluView: Weekly Influenza Surveillance Report. https://www.cdc.gov/flu/weekly/.",
        "Hyndman, R.J. and Athanasopoulos, G. (2021). Forecasting: Principles and Practice. 3rd Ed., OTexts.",
        "Nahmias, S. and Olsen, T.L. (2015). Production and Operations Analysis. 7th Ed., Waveland Press.",
        "Project Proposal: Intelligent Pharmacy Inventory Management. DATA 622 Machine Learning, 2026.",
    ]

    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        run = p.add_run(f"[{i}]  {ref}")
        run.font.size = Pt(10)

    # ==================== SAVE ====================
    output_path = "c:/Projects/DATA622/FinalProjectData622/DATA622_Final_Project_Report.docx"
    doc.save(output_path)
    print(f"Report saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_report()
